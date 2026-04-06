"""Generate PROJECT_PROGRESS.docx from content."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Calibri"

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Global AI Governance Copilot — Project Progress (2-Pager)")
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph("Date: March 2024 | Repo: github.com/avika124/ai-governance-copilot")
doc.add_paragraph()

# Page 1
doc.add_paragraph("Page 1: What Has Been Done").runs[0].bold = True
doc.add_paragraph()

doc.add_paragraph("Data Pipeline (Complete)").runs[0].bold = True
doc.add_paragraph(
    "A Python pipeline ingests EU and India laws, extracts clause-level text, "
    "and stores it in Supabase (PostgreSQL)."
)
doc.add_paragraph("Regulations: 23 (13 EU, 10 India) | Clauses: ~5,147 | Failed: 2 India laws (404)")
doc.add_paragraph()
doc.add_paragraph(
    "EU: GDPR, ePrivacy, Cybersecurity Act, NIS2, AI Act, DSA, DMA, Data Act, "
    "Product Liability, Consumer Rights, DORA, Working Time, Platform Work."
)
doc.add_paragraph(
    "India: DPDP Act, IT Act, BNS, BNSS, Consumer Protection, Competition, RBI, "
    "SEBI, Industrial Relations, Code on Wages, Clinical Establishments."
)
doc.add_paragraph()

doc.add_paragraph("Annotation Prep (Complete)").runs[0].bold = True
doc.add_paragraph(
    "Clauses exported for Label Studio with three labels: risk_type, actor_type, obligation_type."
)
doc.add_paragraph()

doc.add_paragraph("Version Control").runs[0].bold = True
doc.add_paragraph("Project on GitHub; .env excluded; pipeline and annotation prep committed.")
doc.add_paragraph()
doc.add_paragraph()

# Page 2
doc.add_paragraph("Page 2: How We Did It").runs[0].bold = True
doc.add_paragraph()

doc.add_paragraph("Architecture").runs[0].bold = True
doc.add_paragraph("Stack: Python, requests, BeautifulSoup, PyMuPDF, spaCy, psycopg2, Supabase.")
doc.add_paragraph("Flow: config.py → fetch_eu.py / fetch_india.py → extract_clauses.py → db_client.py → Supabase.")
doc.add_paragraph()

doc.add_paragraph("EU Ingestion").runs[0].bold = True
doc.add_paragraph("Source: EUR-Lex HTML. Method: requests + BeautifulSoup; regex splits on \"Article X\". Fallback: PyMuPDF.")
doc.add_paragraph()

doc.add_paragraph("India Ingestion").runs[0].bold = True
doc.add_paragraph(
    "Source: legislative.gov.in PDFs. Method: requests.Session with Referer, homepage visit for cookies. "
    "Parsing: PyMuPDF extracts text; regex splits on \"Section X\"."
)
doc.add_paragraph()

doc.add_paragraph("Clause Extraction").runs[0].bold = True
doc.add_paragraph("Tool: spaCy en_core_web_sm. Rules: Skip <30 chars, numbers, headers; UUID4 per clause; batch insert (50).")
doc.add_paragraph()

doc.add_paragraph("Database (Supabase)").runs[0].bold = True
doc.add_paragraph("regulations: regulation_id, country, law_name, law_category, law_type, year, source_url, raw_text")
doc.add_paragraph("clauses: clause_id, regulation_id, article_number, clause_text, char_count")
doc.add_paragraph()

doc.add_paragraph("Annotation Export").runs[0].bold = True
doc.add_paragraph("python -m annotate.prepare_labelstudio [limit] → annotated/labelstudio_import.json")
doc.add_paragraph("Config: LABELSTUDIO_CONFIG.xml")
doc.add_paragraph()

doc.add_paragraph("Run Commands").runs[0].bold = True
doc.add_paragraph(
    "cd data_pipeline && pip install -r requirements.txt && python -m spacy download en_core_web_sm"
)
doc.add_paragraph("cp .env.example .env  # SUPABASE_DB_URL")
doc.add_paragraph("python run_pipeline.py")
doc.add_paragraph("python -m annotate.prepare_labelstudio 500")
doc.add_paragraph()

doc.add_paragraph("Next Steps (Not Built)").runs[0].bold = True
doc.add_paragraph("Model training (DeBERTa, FAISS) | Agent pipeline (LangChain) | FastAPI | Streamlit UI")

doc.save(Path(__file__).parent / "PROJECT_PROGRESS.docx")
print("Saved PROJECT_PROGRESS.docx")
